"""
ResearchCamp — Multi-Agent Researcher (fixed + optimized)

Extra deps needed for the export buttons:
    pip install python-docx fpdf2

This file is now a thin UI layer: all orchestration (search -> reader ->
writer -> critic) lives in src/pipelines/pipeline.py and is shared with
main.py, instead of being duplicated here.
"""

import base64
import io
import re
import time

import streamlit as st
import streamlit.components.v1 as components

from src.pipelines.pipeline import run_research_pipeline, PipelineStepError

# Optional export libs
from docx import Document
from fpdf import FPDF

# ── Page config ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="ResearchCamp: Multi-Agent Researcher",
    page_icon="",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ── Custom CSS ────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Syne:wght@400;600;700;800&family=DM+Mono:wght@300;400;500&family=DM+Sans:ital,wght@0,300;0,400;0,500;1,300&display=swap');

/* ── Reset & base ── */
html, body, [class*="css"] {
    font-family: 'DM Sans', sans-serif;
    color: #f2f2f2;
}

.stApp {
    background: #060606;
    background-image:
        radial-gradient(circle at top left, rgba(255,255,255,0.06), transparent 32%),
        radial-gradient(circle at bottom right, rgba(255,255,255,0.05), transparent 30%),
        linear-gradient(180deg, #050505 0%, #0d0d0d 100%);
}

/* ── Hide default streamlit chrome ── */
#MainMenu, footer, header { visibility: hidden; }
.block-container { padding: 2rem 3rem 4rem; max-width: 1200px; }

/* ── Hero header ── */
.hero {
    text-align: center;
    padding: 3.5rem 0 2.5rem;
    position: relative;
}
.hero-eyebrow {
    font-family: 'DM Mono', monospace;
    font-size: 0.7rem;
    font-weight: 500;
    letter-spacing: 0.25em;
    text-transform: uppercase;
    color: #cfcfcf;
    margin-bottom: 1rem;
    opacity: 0.85;
}
.hero h1 {
    font-family: 'Syne', sans-serif;
    font-size: clamp(2.8rem, 6vw, 5rem);
    font-weight: 800;
    line-height: 1.0;
    letter-spacing: -0.03em;
    color: #ffffff;
    margin: 0 0 1rem;
}
.hero h1 span {
    background: linear-gradient(135deg, #ffffff, #8a8a8a);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
}
.hero-sub {
    font-size: 1.05rem;
    font-weight: 300;
    color: #b8b8b8;
}

/* ── Divider ── */
.divider {
    height: 1px;
    background: linear-gradient(90deg, transparent, rgba(255,255,255,0.28), transparent);
    margin: 2rem 0;
}

/* ── Input card ── */
.st-key-input_card {
    background: rgba(255,255,255,0.035);
    border: 1px solid rgba(255,255,255,0.14);
    border-radius: 22px;
    padding: 2rem 2.5rem;
    margin-bottom: 2rem;
    backdrop-filter: blur(14px);
    box-shadow: 0 10px 40px rgba(0,0,0,0.5);
}

/* ── Streamlit input overrides ── */
.stTextInput > div > div > input {
    background: rgba(255,255,255,0.05) !important;
    border: 1px solid rgba(255,255,255,0.2) !important;
    border-radius: 12px !important;
    color: #ffffff !important;
    font-family: 'DM Sans', sans-serif !important;
    font-size: 1rem !important;
    padding: 0.8rem 1rem !important;
    transition: all 0.2s ease !important;
}
.stTextInput > div > div > input:focus {
    border-color: #ffffff !important;
    box-shadow: 0 0 0 4px rgba(255,255,255,0.1) !important;
}
.stTextInput > label {
    font-family: 'DM Mono', monospace !important;
    font-size: 0.72rem !important;
    letter-spacing: 0.15em !important;
    text-transform: uppercase !important;
    color: #cfcfcf !important;
    font-weight: 500 !important;
}
.stTextArea > div,
.stTextArea > div > div,
.stTextArea div[data-baseweb="textarea"],
.stTextArea div[data-baseweb="base-input"] {
    background: transparent !important;
    background-color: transparent !important;
    border: none !important;
    box-shadow: none !important;
}

.stTextArea div[data-baseweb="textarea"] {
    border: 1px solid rgba(255,255,255,0.15) !important;
    border-radius: 12px !important;
    transition: border-color 0.2s ease !important;
}

.stTextArea div[data-baseweb="textarea"]:focus-within {
    border-color: rgba(255,255,255,0.5) !important;
}

.stTextArea textarea {
    background: transparent !important;
    color: #ffffff !important;
    font-family: 'DM Sans', sans-serif !important;
    font-size: 1.15rem !important;
    font-weight: 300 !important;
    line-height: 1.6 !important;
    padding: 1rem !important;
    box-shadow: none !important;
    height: 50px;
}

.stTextArea textarea::placeholder {
    color: rgba(255,255,255,0.35) !important;
    font-weight: 300 !important;
}

/* ── Action row (upload + send) pinned to the bottom-right of the card ── */
.st-key-input_card { padding-bottom: 1.2rem !important; }
.st-key-input_actions { margin-top: -0.5rem; }

.st-key-input_actions [data-testid="column"] {
    display: flex;
    align-items: center;
    justify-content: flex-end;
}

/* ── Send button: normal rounded rectangle ── */
.st-key-send_btn .stButton > button {
    width: 200px !important;
    min-height: 44px !important;
    border-radius: 12px !important;
    font-size: 2rem !important;
    font-weight: 700 !important;
    align-items: center;
    justify-content: center;
    box-shadow: 0 6px 20px rgba(255,255,255,0.18) !important;
    background: linear-gradient(135deg, #ffffff 0%, #9a9a9a 100%) !important;
    color: #0a0a0a !important;
}

/* ── File upload: collapse the big dropzone into a small "+" circle ── */
.st-key-file_upload {
    display: flex;
    justify-content: center;
}

.st-key-file_upload [data-testid="stFileUploaderDropzone"] {
    background: rgba(255,255,255,0.06) !important;
    border: 1px solid rgba(255,255,255,0.18) !important;
    border-radius: 50% !important;
    width: 44px !important;
    height: 44px !important;
    min-height: 44px !important;
    padding: 0 !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
    position: relative;
    overflow: hidden;
}

.st-key-file_upload [data-testid="stFileUploaderDropzone"]::before {
    content: "+";
    font-size: 1.6rem;
    color: #ffffff;
    position: absolute;
    pointer-events: none;
}

.st-key-file_upload [data-testid="stFileUploaderDropzoneInstructions"],
.st-key-file_upload [data-testid="stFileUploaderDropzoneIcon"],
.st-key-file_upload svg,
.st-key-file_upload small,
.st-key-file_upload button {
    display: none !important;
}

.st-key-file_upload input[type="file"] {
    position: absolute !important;
    width: 100% !important;
    height: 100% !important;
    cursor: pointer !important;
    opacity: 0 !important;
    z-index: 10 !important;
}

.st-key-file_upload [data-testid="stFileUploaderFileData"] {
    display: none !important;
}

/* ── Buttons (primary "Build Report" + download buttons) ── */
.stButton > button, .stDownloadButton > button {
    background: linear-gradient(135deg, #ffffff 0%, #9a9a9a 100%) !important;
    color: #0a0a0a !important;
    font-family: 'Syne', sans-serif !important;
    font-weight: 700 !important;
    font-size: 0.9rem !important;
    letter-spacing: 0.04em !important;
    border: none !important;
    border-radius: 12px !important;
    padding: 0.75rem 1.6rem !important;
    cursor: pointer !important;
    transition: all 0.18s ease !important;
    box-shadow: 0 8px 30px rgba(255,255,255,0.12) !important;
    width: 100%;
}
.stButton > button:hover, .stDownloadButton > button:hover {
    transform: translateY(-2px) scale(1.01) !important;
    box-shadow: 0 12px 35px rgba(255,255,255,0.22) !important;
}
.stButton > button:active, .stDownloadButton > button:active {
    transform: translateY(0) !important;
}

/* ── Pipeline step cards ── */
.step-card {
    background: rgba(255,255,255,0.03);
    border: 1px solid rgba(255,255,255,0.09);
    border-radius: 18px;
    padding: 1.5rem 1.8rem;
    margin-bottom: 1.2rem;
    position: relative;
    overflow: hidden;
    transition: all 0.25s ease;
    backdrop-filter: blur(10px);
}
.step-card:hover {
    transform: translateY(-2px);
    border-color: rgba(255,255,255,0.2);
}
.step-card.working {
    border-color: rgba(255,157,66,0.55);
    background: rgba(255,157,66,0.06);
}
.step-card.done {
    border-color: rgba(200,200,200,0.35);
    background: rgba(255,255,255,0.04);
}
.step-card.error {
    border-color: rgba(255,90,90,0.6);
    background: rgba(255,90,90,0.08);
}
.step-card::before {
    content: '';
    position: absolute;
    left: 0; top: 0; bottom: 0;
    width: 4px;
    border-radius: 18px 0 0 18px;
    background: rgba(255,255,255,0.08);
    transition: background 0.3s;
}
.step-card.working::before { background: #ff9d42; }
.step-card.done::before    { background: #c9c9c9; }
.step-card.error::before   { background: #ff5a5a; }

.step-header {
    display: flex;
    align-items: center;
    gap: 0.8rem;
    margin-bottom: 0.3rem;
}
.step-num {
    font-family: 'DM Mono', monospace;
    font-size: 0.68rem;
    font-weight: 500;
    letter-spacing: 0.15em;
    color: #cfcfcf;
    opacity: 0.85;
}
.step-title {
    font-family: 'Syne', sans-serif;
    font-size: 1rem;
    font-weight: 700;
    color: #ffffff;
}
.step-status {
    margin-left: auto;
    font-family: 'DM Mono', monospace;
    font-size: 0.68rem;
    letter-spacing: 0.1em;
}
.status-waiting  { color: #6b6b6b; }
.status-working  { color: #ff9d42; }
.status-done     { color: #c9c9c9; }
.status-error    { color: #ff5a5a; }

/* ── Result panels (raw search/reader output) ── */
.result-panel {
    background: rgba(255,255,255,0.025);
    border: 1px solid rgba(255,255,255,0.09);
    border-radius: 18px;
    padding: 1.8rem 2rem;
    margin-top: 1rem;
    margin-bottom: 1.5rem;
    backdrop-filter: blur(12px);
}
.result-panel-title {
    font-family: 'DM Mono', monospace;
    font-size: 0.7rem;
    font-weight: 500;
    letter-spacing: 0.2em;
    text-transform: uppercase;
    color: #cfcfcf;
    margin-bottom: 1rem;
    padding-bottom: 0.7rem;
    border-bottom: 1px solid rgba(255,255,255,0.14);
}
.result-content {
    font-size: 0.92rem;
    line-height: 1.8;
    color: #dcdcdc;
    white-space: pre-wrap;
    font-family: 'DM Sans', sans-serif;
}

/* ── Report & feedback panels ── */
.st-key-report_panel {
    background: rgba(255,255,255,0.025);
    border: 1px solid rgba(255,255,255,0.18);
    border-radius: 20px;
    padding: 2rem 2.5rem;
    margin-top: 1rem;
    backdrop-filter: blur(14px);
}
.st-key-feedback_panel {
    background: rgba(255,255,255,0.025);
    border: 1px solid rgba(200,200,200,0.18);
    border-radius: 20px;
    padding: 2rem 2.5rem;
    margin-top: 1rem;
    backdrop-filter: blur(14px);
}
.panel-label {
    font-family: 'DM Mono', monospace;
    font-size: 0.7rem;
    letter-spacing: 0.2em;
    text-transform: uppercase;
    margin-bottom: 1.2rem;
    padding-bottom: 0.7rem;
}
.panel-label.orange {
    color: #ffffff;
    border-bottom: 1px solid rgba(255,255,255,0.14);
}
.panel-label.green {
    color: #c9c9c9;
    border-bottom: 1px solid rgba(200,200,200,0.14);
}

/* ── Expander ── */
details {
    background: rgba(255,255,255,0.015);
    border-radius: 14px;
    padding: 0.3rem 0.8rem;
    border: 1px solid rgba(255,255,255,0.06);
}
details summary {
    font-family: 'DM Mono', monospace !important;
    font-size: 0.75rem !important;
    color: #b8b8b8 !important;
    letter-spacing: 0.1em !important;
    cursor: pointer;
}

/* ── Section heading ── */
.section-heading {
    font-family: 'Syne', sans-serif;
    font-size: 1.35rem;
    font-weight: 700;
    color: #ffffff;
    margin: 2rem 0 1rem;
}

/* ── Toast-style notice ── */
.notice {
    font-family: 'DM Mono', monospace;
    font-size: 0.72rem;
    color: #7a7a7a;
    text-align: center;
    margin-top: 3rem;
    letter-spacing: 0.08em;
}
</style>
""", unsafe_allow_html=True)


STEPS = [
    ("search", "01", "Search Agent", "Gathers recent web information"),
    ("reader", "02", "Reader Agent", "Scrapes & extracts deep content"),
    ("writer", "03", "Writer Chain", "Drafts the full research report"),
    ("critic", "04", "Critic Chain", "Reviews & scores the report"),
]
STEP_LOOKUP = {k: (n, t, d) for k, n, t, d in STEPS}


# ── Helper: render a step card into a placeholder ─────────────────────────────
def render_step(placeholder, num: str, title: str, desc: str, state: str, error_msg: str = ""):
    """Statuses are 'waiting' / 'working' / 'done' / 'error'."""
    status_map = {
        "waiting": ("WAITING", "status-waiting", ""),
        "working": ("● WORKING", "status-working", "working"),
        "done":    ("✓ DONE",   "status-done",    "done"),
        "error":   ("✕ FAILED", "status-error",   "error"),
    }
    label, status_cls, card_cls = status_map.get(state, status_map["waiting"])

    extra = ""
    if state == "error" and error_msg:
        extra = f"<div style='font-size:0.78rem;color:#ff8a8a;margin-top:0.4rem;'>{error_msg}</div>"
    elif desc:
        extra = f"<div style='font-size:0.82rem;color:#9a9a9a;margin-top:0.3rem;'>{desc}</div>"

    placeholder.markdown(f"""
    <div class="step-card {card_cls}">
        <div class="step-header">
            <span class="step-num">{num}</span>
            <span class="step-title">{title}</span>
            <span class="step-status {status_cls}">{label}</span>
        </div>
        {extra}
    </div>
    """, unsafe_allow_html=True)


# ── Helpers: export the final report ──────────────────────────────────────────
def markdown_to_docx_bytes(md_text: str) -> bytes:
    doc = Document()
    for raw_line in md_text.split("\n"):
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)", line)
        bullet = re.match(r"^[-*]\s+(.*)", line)
        numbered = re.match(r"^\d+\.\s+(.*)", line)
        if heading:
            level = min(len(heading.group(1)), 4)
            doc.add_heading(heading.group(2), level=level)
        elif bullet:
            doc.add_paragraph(bullet.group(1), style="List Bullet")
        elif numbered:
            doc.add_paragraph(numbered.group(1), style="List Number")
        else:
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def markdown_to_pdf_bytes(md_text: str) -> bytes:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)
    for raw_line in md_text.split("\n"):
        line = raw_line.strip()
        if not line:
            pdf.ln(4)
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)", line)
        if heading:
            size = max(18 - 2 * (len(heading.group(1)) - 1), 12)
            text = heading.group(2).encode("latin-1", "replace").decode("latin-1")
            pdf.set_font("Helvetica", "B", size)
            pdf.multi_cell(0, 8, text)
            pdf.set_font("Helvetica", size=11)
        else:
            safe = line.encode("latin-1", "replace").decode("latin-1")
            pdf.multi_cell(0, 6, safe)
    output = pdf.output()
    return bytes(output)


def copy_to_clipboard_button(text: str, key: str):
    b64 = base64.b64encode(text.encode("utf-8")).decode("utf-8")
    components.html(f"""
    <div style="font-family:'Syne',sans-serif;">
        <button id="copyBtn_{key}" style="
            background:linear-gradient(135deg,#ffffff 0%, #9a9a9a 100%);
            color:#0a0a0a;
            font-family:'Syne',sans-serif;
            font-weight:700;
            font-size:0.9rem;
            letter-spacing:0.04em;
            border:none;
            border-radius:12px;
            padding:0.75rem 1.6rem;
            cursor:pointer;
            width:100%;
            box-shadow:0 8px 30px rgba(255,255,255,0.12);
        ">⧉ Copy Report</button>
    </div>
    <script>
        const btn_{key} = document.getElementById("copyBtn_{key}");
        btn_{key}.addEventListener("click", function() {{
            const decoded = decodeURIComponent(escape(window.atob("{b64}")));
            navigator.clipboard.writeText(decoded).then(function() {{
                btn_{key}.innerText = "✓ Copied!";
                setTimeout(function() {{ btn_{key}.innerText = "⧉ Copy Report"; }}, 1500);
            }});
        }});
    </script>
    """, height=60)


# ── Session state init ────────────────────────────────────────────────────────
for key, default in (("results", {}), ("running", False), ("done", False), ("error", None)):
    if key not in st.session_state:
        st.session_state[key] = default


# ── Hero ──────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="hero">
    <div class="hero-eyebrow">Multi-Agent AI System</div>
    <h1>Research<span>Camp</span></h1>
    <p class="hero-sub">
       Did your dog eat your homework ? build it in few minutes with researchcamp .
    </p>
<div class="divider"></div>
</div>
""", unsafe_allow_html=True)


with st.container(key="input_card"):
    research_request = st.text_area(
        "Research Request",
        placeholder="e.g. Research the roadmap for AGI development over the next 5 years...",
        key="research_input",
        label_visibility="collapsed",
        height=120,
    )

    with st.container(key="input_actions"):
        spacer_col, upload_col, send_col = st.columns([7, 1, 2])
        with upload_col:
            uploaded_file = st.file_uploader(
                "Attach file",
                key="file_upload",
                label_visibility="collapsed",
            )
        with send_col:
            run_btn = st.button("build report ", key="send_btn", use_container_width=True)


# ── Kick off a run ────────────────────────────────────────────────────────────
if run_btn:
    if not research_request.strip():
        st.warning("Please enter a research topic first.")
    else:
        st.session_state.results = {}
        st.session_state.running = True
        st.session_state.done = False
        st.session_state.error = None
        st.rerun()


# ── Layout: input left, pipeline right ───────────────────────────────────────
col_input, col_spacer, col_pipeline = st.columns([5, 0.5, 4])

with col_pipeline:
    st.markdown('<div class="section-heading">Pipeline</div>', unsafe_allow_html=True)

    step_placeholders = {}
    for step_key, num, title, desc in STEPS:
        step_placeholders[step_key] = st.empty()
        if step_key in st.session_state.results:
            initial_state = "done"
        elif st.session_state.error and st.session_state.error["step"] == step_key:
            initial_state = "error"
        else:
            initial_state = "waiting"
        err_msg = st.session_state.error["message"] if (
            st.session_state.error and st.session_state.error["step"] == step_key
        ) else ""
        render_step(step_placeholders[step_key], num, title, desc, initial_state, err_msg)


# ── Run pipeline (updates the pipeline cards live via callbacks) ─────────────
if st.session_state.running and not st.session_state.done:

    topic_val = st.session_state.research_input

    def on_step_start(step_key: str):
        num, title, desc = STEP_LOOKUP[step_key]
        render_step(step_placeholders[step_key], num, title, desc, "working")

    def on_step_done(step_key: str, result):
        num, title, desc = STEP_LOOKUP[step_key]
        render_step(step_placeholders[step_key], num, title, desc, "done")
        st.session_state.results = dict(st.session_state.results, **{step_key: result})

    try:
        run_research_pipeline(topic_val, on_step_start=on_step_start, on_step_done=on_step_done)
        st.session_state.error = None
    except PipelineStepError as e:
        # A step failed even after internal retries (e.g. the model kept
        # producing malformed tool calls). Mark that step's card as failed
        # instead of letting the whole app crash, and keep whatever earlier
        # steps already succeeded so the user isn't left with nothing.
        num, title, desc = STEP_LOOKUP[e.step]
        render_step(step_placeholders[e.step], num, title, desc, "error", str(e.original))
        st.session_state.error = {"step": e.step, "message": str(e.original)}

    st.session_state.running = False
    st.session_state.done = True
    st.rerun()


# ── Results display (always reads FRESH state, never a stale snapshot) ───────
r = st.session_state.results

with col_input:
    if st.session_state.error:
        st.error(
            f"The **{st.session_state.error['step']}** step failed after retries: "
            f"{st.session_state.error['message']}"
        )

    if r:
        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
        st.markdown('<div class="section-heading">Results</div>', unsafe_allow_html=True)

        if "search" in r:
            with st.expander(" Search Results (raw)", expanded=False):
                st.markdown(f"""
                    <div class="result-panel">
                        <div class="result-panel-title">Search Agent Output</div>
                        <div class="result-content">{r["search"]}</div>
                    </div>
                    """, unsafe_allow_html=True)

        if "reader" in r:
            with st.expander(" Scraped Content (raw)", expanded=False):
                st.markdown(f"""
                    <div class="result-panel">
                        <div class="result-panel-title">Reader Agent Output</div>
                        <div class="result-content">{r["reader"]}</div>
                    </div>
                    """, unsafe_allow_html=True)

    elif not st.session_state.error:

        st.markdown("""
        <div style="display:flex;gap:0.5rem;flex-wrap:wrap;margin-bottom:1.5rem;">
            <span style="font-family:'DM Mono',monospace;font-size:0.68rem;color:#7a7a7a;letter-spacing:0.1em;">
                TRY →
            </span>
        """, unsafe_allow_html=True)

        examples = [
            "Future of LLM in Tech Industry",
            "All Lastest AI Agents in 2026",
            "Roadmap for AGI development in next 5 years",
        ]
        for ex in examples:
            st.markdown(f"""
            <span style="
                background:rgba(255,255,255,0.04);
                border:1px solid rgba(255,255,255,0.1);
                border-radius:8px;
                padding:0.35rem 0.8rem;
                font-size:0.75rem;
                color:#dcdcdc;
                font-family:'DM Sans',sans-serif;
                cursor:default;
            ">
                {ex}
            </span>
            """, unsafe_allow_html=True)

        st.markdown("</div>", unsafe_allow_html=True)

if "writer" in r:
        with st.container(key="report_panel"):
            st.markdown('<div class="panel-label orange">Final Research Report</div>', unsafe_allow_html=True)
            st.markdown(r["writer"])

            st.markdown("<div style='height:0.5rem;'></div>", unsafe_allow_html=True)
            btn_cols = st.columns(4)
            with btn_cols[0]:
                st.download_button(
                    "⬇ .md",
                    data=r["writer"],
                    file_name=f"research_report_{int(time.time())}.md",
                    mime="text/markdown",
                    use_container_width=True,
                )
            with btn_cols[1]:
                st.download_button(
                    "⬇ .docx",
                    data=markdown_to_docx_bytes(r["writer"]),
                    file_name=f"research_report_{int(time.time())}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )

            with btn_cols[3]:
                copy_to_clipboard_button(r["writer"], key="report")

if "critic" in r:
        with st.container(key="feedback_panel"):
            st.markdown('<div class="panel-label green">Critic Feedback</div>', unsafe_allow_html=True)
            st.markdown(r["critic"])


# ── Footer ────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="notice">
    ResearchAgent · Powered by LangChain multi-agent pipeline · Built with Streamlit
</div>
""", unsafe_allow_html=True)